from PyQt5.QtWidgets import QWidget, QVBoxLayout, QTextEdit, QPushButton, QLabel, QFileDialog, QMessageBox, QApplication
from PyQt5.QtGui import QFont
import json
from docx import Document
from datetime import datetime
import os

# Define the path for the template
template_path = os.path.join(os.path.dirname(__file__), "../templates/Agenda_Blank_With_Styles.docx")

def load_ui(parent=None):
    """Set up and return the UI for the agenda generator."""
    widget = QWidget(parent)
    layout = QVBoxLayout(widget)

    # JSON input text box
    json_textbox = QTextEdit()
    json_textbox.setPlaceholderText("Paste JSON Here")
    json_textbox.setFont(QFont("Courier", 10))
    layout.addWidget(json_textbox)

    # Paste from clipboard button
    paste_button = QPushButton("Paste from Clipboard")
    paste_button.clicked.connect(lambda: json_textbox.setPlainText(QApplication.clipboard().text()))
    layout.addWidget(paste_button)

    # Generate Agenda button
    generate_button = QPushButton("Generate Agenda")
    generate_button.clicked.connect(lambda: generate_agenda(json_textbox.toPlainText(), widget))
    layout.addWidget(generate_button)

    # Status label
    status_label = QLabel()
    layout.addWidget(status_label)

    # Define function to handle agenda generation
    def generate_agenda(json_data, parent):
        try:
            # Validate JSON format
            data = json.loads(json_data)

            # Populate the template
            doc, output_filename = populate_template(data)

            # Prompt user to select save location
            filename, _ = QFileDialog.getSaveFileName(
                parent,
                "Save Agenda",
                output_filename,
                "Word Documents (*.docx)"
            )

            if filename:
                # Save document to the selected location
                doc.save(filename)
                status_label.setText(f"Agenda saved to {filename}")
                QMessageBox.information(parent, "Success", f"Agenda saved successfully to {filename}")
            else:
                status_label.setText("Save cancelled.")
        except json.JSONDecodeError:
            QMessageBox.warning(parent, "Error", "Invalid JSON format.")
        except ValueError as e:
            status_label.setText(str(e))
            QMessageBox.warning(parent, "Error", str(e))
        except Exception as e:
            status_label.setText(f"Error: {e}")
            QMessageBox.critical(parent, "Error", f"An error occurred: {e}")

    return widget

def populate_template(data):
    """Generate the agenda document from validated JSON data."""
    doc = Document(template_path)
    
    # Extract client info
    client_name = data.get("client", {}).get("name", "Unknown Client")
    client_date = datetime.strptime(data.get("client", {}).get("date", "Unknown Date"), "%B %d, %Y")
    doc.add_paragraph(f"Agenda: {client_name}", style="TitleAgenda")
    doc.add_paragraph(client_date.strftime("%B %d, %Y"), style="DateAgenda")

    # Add summary section
    summary = data.get("summary", {})
    total_value = summary.get("total_value", "$0")
    total_income = summary.get("total_income", "$0")
    doc.add_paragraph(f"Review Accounts\tTotal Value: {total_value} Income: {total_income}", style="List_Category")

    # Add account details
    for account in data.get("accounts", []):
        account_count = account.get("count", "N/A")
        last_four = account.get("last_four", "N/A")
        doc.add_paragraph(f"Account {account_count} xxxx-{last_four}", style="List_SubCategory")
        
        account_value = account.get("account_value", "$0")
        account_cash_flow = account.get("account_cash_flow", "$0")
        account_performance_ytd = account.get("account_performance_ytd", "N/A")
        account_allocation = account.get("account_allocation", "N/A")
        
        doc.add_paragraph(f"Total Account:\t{account_value}", style="List_Item")
        doc.add_paragraph(f"Current Cash Flow:\t{account_cash_flow}", style="List_Item")
        doc.add_paragraph(f"Performance YTD:\t{account_performance_ytd}", style="List_Item")
        doc.add_paragraph(f"Allocation:\t{account_allocation}", style="List_Item")
    
    # Generate output filename based on client name and date
    safe_name = "".join(c for c in client_name if c.isalnum() or c in (' ', '-'))
    output_filename = f"Agenda_{safe_name}_{client_date.strftime('%Y%m%d')}.docx"
    
    return doc, output_filename

def parse_date(date_str):
    """Parse various date formats to a datetime object."""
    try:
        return datetime.strptime(date_str.replace("st", "").replace("nd", "").replace("rd", "").replace("th", ""), "%B %d, %Y")
    except ValueError:
        try:
            return datetime.strptime(date_str, "%m/%d/%Y")
        except ValueError:
            try:
                return datetime.strptime(date_str, "%Y-%m-%d")
            except ValueError:
                raise ValueError(f"Invalid date format: {date_str}")

def format_dollar_amount(amount):
    """Format dollar amounts consistently."""
    if isinstance(amount, str):
        amount = amount.replace("$", "").replace(",", "")
        try:
            amount = float(amount)
        except ValueError:
            raise ValueError(f"Invalid dollar amount: {amount}")
    return "${:,.0f}".format(amount)
